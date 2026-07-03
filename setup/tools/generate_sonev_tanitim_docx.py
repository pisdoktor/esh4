#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""SONEV tanıtım belgesi (.docx) üretici — docs/SONEV_TANITIM.docx"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SONEV_TANITIM.docx"


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)

    for level, size in [(1, 16), (2, 13), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        h.font.size = Pt(size)
        h.paragraph_format.space_before = Pt(12 if level > 1 else 0)
        h.paragraph_format.space_after = Pt(6)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SONEV\nEvde Sağlık Hizmetleri Bilgi Yönetim Sistemi")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Kapsamlı Tanıtım ve Karşılaştırmalı Değerlendirme Belgesi")
    r.font.size = Pt(14)
    r.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\nSürüm 4.0.0 (Burutay)\nBelge tarihi: {date.today().strftime('%d.%m.%Y')}\n").font.size = Pt(11)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Bu belge A4 baskıya uygun metin formatında hazırlanmıştır.\n"
        "Görsel öğelerden ziyade okunabilirlik ve içerik bütünlüğü önceliklidir."
    )
    nr.font.size = Pt(10)
    nr.italic = True
    nr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()


def add_paragraphs(doc: Document, text: str) -> None:
    for block in text.strip().split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("• "):
            for line in block.split("\n"):
                line = line.strip()
                if line.startswith("• "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)
        else:
            doc.add_paragraph(block)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    set_document_defaults(doc)
    add_title_page(doc)

    # --- 1. GİRİŞ ---
    doc.add_heading("1. Giriş ve Belgenin Amacı", level=1)
    add_paragraphs(
        doc,
        """
Bu belge, SONEV (ESH — Evde Sağlık Hizmetleri) bilgi yönetim sisteminin kapsamlı tanıtımını sunmak amacıyla hazırlanmıştır. Belge; sistemin işlevsel kapsamını, teknik mimarisini, hedef kullanıcı profillerini ve operasyonel iş akışlarını ayrıntılı biçimde açıklar. Bunun yanında SONEV'in Türkiye Cumhuriyeti Sağlık Bakanlığı'nın resmî Evde Sağlık Yönetim Sistemi (ESYS) ile ilişkisini, ESYS'e kıyasla sunduğu operasyonel ve analitik üstünlükleri ve dünya genelindeki benzer evde bakım / saha sağlık yazılımlarıyla karşılaştırmalı değerlendirmesini içerir.

SONEV, evde sağlık birimlerinin günlük operasyonlarını yönetmek için geliştirilmiş kurumsal bir web uygulamasıdır. Hasta kaydından günlük ziyaret planlamasına, saha izlemlerinden istatistik ve raporlamaya, personel ve stok yönetiminden kurumsal güvenlik ve uyum süreçlerine kadar geniş bir işlev yelpazesini tek platformda birleştirir. Sistem PHP 8 tabanlı özel bir MVC mimarisi üzerinde çalışır; varsayılan veritabanı MariaDB/MySQL olmakla birlikte Microsoft SQL Server, PostgreSQL, Oracle ve SQLite şemaları da desteklenir.

Belge, kurum içi bilgilendirme, yatırım kararı desteği, teknik değerlendirme ve dış paydaşlara sistem tanıtımı gibi amaçlarla kullanılabilir. Metin, slayt sunumlarından farklı olarak A4 kağıda basıldığında rahat okunacak paragraf yapısında düzenlenmiştir.
        """,
    )

    # --- 2. BAĞLAM ---
    doc.add_heading("2. Evde Sağlık Hizmetleri ve Dijitalleşme Bağlamı", level=1)
    add_paragraphs(
        doc,
        """
Evde sağlık hizmetleri (ESH), hastanın konutunda sunulan; muayene, tetkik, tedavi, tıbbi bakım, rehabilitasyon, eğitim ve danışmanlık gibi hizmetleri kapsayan karmaşık bir sağlık modelidir. Bu hizmet modeli, özellikle kronik hastalığı olan, hareket kısıtlılığı bulunan, palyatif bakım ihtiyacı olan veya taburcu sonrası evde takip gerektiren hastalar için kritik öneme sahiptir.

Türkiye'de evde sağlık hizmetleri, Sağlık Bakanlığı koordinasyonunda yürütülmekte; başvurular e-Nabız, ESHİM (444 38 33) ve ESYS üzerinden alınmakta; Evde Sağlık Hizmetleri Koordinasyon Merkezi (ESKOM) ile aile hekimliği ve hastane bilgi sistemleri arasında dijital entegrasyon hedeflenmektedir. Resmî yönetmelikler, ESYS'in AHBS (Aile Hekimliği Bilgi Sistemi), HBYS (Hastane Bilgi Yönetim Sistemi) ve UHDS (Uzaktan Hasta Değerlendirme Sistemi) ile entegre çalışan ulusal dijital sağlık altyapısının merkezinde yer aldığını tanımlar.

Bu ulusal çerçevede ESYS, başvuru, değerlendirme, izlem, sevk, nakil ve taburculuk verilerinin Bakanlık standartlarına uygun kayıt altına alınması için zorunlu ve merkezî bir sistemdir. Ancak saha birimlerinin günlük operasyonel ihtiyaçları — günlük rota planlaması, ekip koordinasyonu, detaylı klinik profil takibi, yerel stok yönetimi, kapsamlı istatistik, personel nöbet planı, SMS bildirimi, iç mesajlaşma — çoğu zaman birim düzeyinde daha derin ve esnek araçlar gerektirir. SONEV tam bu boşluğu doldurmak üzere tasarlanmıştır: ulusal sisteme uyumlu kalırken birim operasyonlarını güçlendiren, yerel veri egemenliği sunan ve ESYS ile köprülenebilen bir operasyonel platform.
        """,
    )

    # --- 3. SONEV ÖZET ---
    doc.add_heading("3. SONEV Sisteminin Genel Tanımı", level=1)
    add_paragraphs(
        doc,
        """
SONEV (kod adı: Burutay, sürüm 4.0.0), Türkiye'deki evde sağlık birimlerinin ihtiyaçları gözetilerek geliştirilmiş, modüler ve çok kurumlu (multi-tenant) bir bilgi yönetim sistemidir. Sistem adı kurumsal ayarlardan özelleştirilebilir; varsayılan kısa ad «SONEV» olarak kullanılır.

Sistemin temel tasarım ilkeleri şunlardır:

• Operasyonel bütünlük: Hasta dosyası, plan, izlem ve raporlama tek veri modeli üzerinde birleşir.
• Modülerlik: İsteğe bağlı modüller ayar panelinden açılıp kapatılabilir; kurum ihtiyacına göre özelleştirme yapılır.
• Kurum izolasyonu: Çoklu kurum yapısında her birimin verisi mantıksal olarak ayrılır.
• Uyum hazırlığı: ESYS ve USBS/e-Nabız alan eşlemesi ve köprü altyapısı mevcuttur.
• Güvenlik: CSRF koruması, rol tabanlı yetkilendirme, hasta düzeyinde erişim kontrolü, denetim günlüğü ve KVKK bilgilendirmeleri yerleşiktir.
• Bağımsızlık: Composer ve npm çalışma zamanı için zorunlu değildir; kurumun mevcut PHP + MariaDB altyapısında çalışabilir.
• Ölçeklenebilirlik: Tek veritabanı üzerinde birden fazla evde sağlık birimi bağımsız çalışabilir; federasyon modülü ile bölgeler arası veri köprüsü desteklenir.

SONEV yalnızca hasta listesi tutan bir kayıt uygulaması değildir. Sistem; klinik ölçekler (Barthel, Braden, Harizmi, ITAKI, MNA), pansuman planı, yara fotoğrafları, klinik bayraklar (trakeostomi, sonda, izolasyon, bası yarası vb.), klinik karar desteği, günlük rota optimizasyonu, GPS check-in, stok takibi, nöbet planı, SMS ve mesajlaşma gibi saha ve idari süreçleri entegre biçimde yönetir.
        """,
    )

    # --- 4. KULLANICILAR ---
    doc.add_heading("4. Hedef Kullanıcılar ve Yetkilendirme Modeli", level=1)
    add_paragraphs(
        doc,
        """
SONEV farklı yetki seviyelerindeki kullanıcı gruplarına hizmet verecek şekilde tasarlanmıştır.

Saha personeli (hemşire, pratisyen hekim, uzman tabip, sağlık teknikeri, gerontolog, eczacı, sekreter ve diğer unvanlar) günlük işlerinde hasta kartına erişir, izlem girişi yapar, planlı ziyaretleri görüntüler ve günlük plan ekranlarını kullanır. Birim yöneticileri personel koordinasyonu, randevu takibi, planlama, stok ve raporlama gibi idari işlemleri yürütür. Kurum yöneticileri kendi kurumlarına ait tanımları (branş, işlem, istek türleri vb.), kurumsal ayarları ve personel atamalarını yönetir. Süper yöneticiler platform düzeyinde tüm kurumları görebilir, modül ayarlarını yapılandırır, veritabanı bakımı ve ulusal sistem köprüleri gibi ileri düzey işlemleri gerçekleştirir.

Hasta ve bakım verenler, isteğe bağlı hasta portalı üzerinden TC kimlik numarası ve telefon ile oturum açarak plan ve ziyaret özetlerini görüntüleyebilir, SMS onayı gibi işlemleri yapabilir.

Yetkilendirme modeli üç ana seviyeye dayanır: personel (isadmin=0), yönetici (isadmin=1) ve süper yönetici (isadmin=2). Bunun yanında modül bazlı izinler ve hasta düzeyinde erişim kontrolleri (PatientAccessHelper) uygulanır. Kurum izolasyonu (TenantContext, TenantSqlHelper) sayesinde bir kurumun personeli başka kurumun verisini göremez. Süper yöneticiler üst menüdeki kurum filtresiyle belirli bir kuruma odaklanabilir.
        """,
    )

    # --- 5. HASTA YAŞAM DÖNGÜSÜ ---
    doc.add_heading("5. Hasta Yaşam Döngüsü ve Klinik Derinlik", level=1)

    doc.add_heading("5.1. Dosya Durumları ve İş Akışı", level=2)
    add_paragraphs(
        doc,
        """
Hasta dosyası SONEV'de belirli bir iş akışı içinde ilerler. Yeni başvurular önce bekleyen kayıt havuzuna alınır. Gerekli bilgiler tamamlandıktan ve onay süreci geçildikten sonra hasta aktif statüye geçer. Aktif hastalar düzenli izlem ve planlı ziyaret süreçlerine dahil edilir. Hizmetin sona ermesi, nakil veya vefat gibi durumlarda hasta sırasıyla pasif, nakil veya vefat statülerine taşınır. Silinen kayıtlar ayrı bir listede tutulur ve gerektiğinde denetim amacıyla izlenebilir.

Her hasta kartında kimlik ve iletişim bilgileri (TC kimlik, ad-soyad, doğum tarihi, cinsiyet, anne-baba adı, telefon, aile hekimi, kan grubu), adres (ilçe, mahalle, sokak, kapı numarası hiyerarşisi), güvence türü, tanılar (ICD-10 kataloğu), klinik bayraklar ve cihaz bilgileri yer alır. Bakım veren iletişim bilgileri, alerji ve acil klinik not alanları hasta güvenliği açısından öne çıkarılır.
        """,
    )

    doc.add_heading("5.2. Klinik Ölçekler ve Değerlendirme Araçları", level=2)
    add_paragraphs(
        doc,
        """
SONEV, evde sağlık hastalarının fonksiyonel ve klinik durumunu objektif biçimde izlemek için uluslararası kabul görmüş ölçekleri destekler:

• Barthel İndeksi — günlük yaşam aktivitelerinde bağımlılık düzeyi
• Braden Ölçeği — bası yarası riski
• Harizmi Düşme Riski Ölçeği — düşme riski değerlendirmesi
• ITAKI Düşme Riski Ölçeği — alternatif düşme riski ölçeği
• MNA (Mini Nutritional Assessment) — beslenme durumu

Bu ölçeklerin geçmiş değerlendirmeleri hasta kartında saklanır; istatistik modülünde dağılım ve trend analizleri yapılabilir. Klinik karar desteği servisi (ClinicalDecisionSupportService), yüksek riskli ve izlem gecikmesi olan hastaları otomatik olarak tespit ederek dashboard ve istatistik ekranlarında öne çıkarır.
        """,
    )

    doc.add_heading("5.3. Klinik Bayraklar ve Cihaz Profili", level=2)
    add_paragraphs(
        doc,
        """
SONEV, hastanın evde bakım profilini hızlıca anlamak için kapsamlı klinik bayrak sistemi sunar. PatientClinicalFlagsHelper tek kaynak olarak form, görünüm ve istatistik katmanlarında tutarlılık sağlar. Desteklenen bayraklar arasında oksijen bağımlılığı, CPAP/BiPAP, ventilatör, trakeostomi, aspirasyon ihtiyacı, NG/PEG beslenme tüpleri, kolostomi/ileostomi/ürostomi, mesane sondası, port kateter, PICC/CVC, aktif IV tedavi, ev diyalizi, aktif dren, bası yarası, hasta yatağı, izolasyon önlemi, bez ve mama kullanımı yer alır.

Bu bayraklar hasta listelerinde kısa rozetler (O₂, Trakeo, Sonda, İzolasyon vb.) olarak gösterilir; hasta kartında cihaz özeti ve klinik uyarı kutuları oluşturulur. Genel sekmede aktif dosyalarda bazı bayraklar tıklanarak hızlı AJAX güncellemesi yapılabilir. Bu yaklaşım, saha ekibinin hasta kartını açmadan önce kritik bilgileri görmesini sağlar.
        """,
    )

    doc.add_heading("5.4. Adres, Harita ve Rota", level=2)
    add_paragraphs(
        doc,
        """
Adres yönetimi evde sağlık planlamasının temelidir. SONEV'de adres verileri ilçe → mahalle → sokak → kapı numarası hiyerarşisinde tutulur. Belediye adres senkronu (Denizli modeli) ile dış kaynaklardan adres ağacı güncellenebilir. Koordinat bilgisi olan hastalar harita üzerinde gösterilir.

Dashboard ve rota modülleri TomTom veya alternatif harita sağlayıcıları ile entegre çalışır. Günlük ziyaret planı için rota optimizasyonu, mesafe matrisi hesaplama ve harita üzerinde görselleştirme desteklenir. İzlem kayıtlarında GPS check-in (saha konum kaydı) özelliği bulunur; yapıldı olarak işaretlenen izlemlerde GPS konumu zorunlu kılınabilir. Bu özellik, saha doğrulaması ve operasyonel şeffaflık açısından ESYS'in standart izlem kayıtlarının ötesinde bir katman sunar.
        """,
    )

    # --- 6. MODÜLLER ---
    doc.add_heading("6. Modül Yapısı ve İşlevsel Kapsam", level=1)

    doc.add_heading("6.1. Çekirdek Modüller", level=2)
    add_paragraphs(
        doc,
        """
Ana panel (Dashboard): Günlük ziyaret planı, takvim görünümü, klinik karar desteği özeti, rota ekranları ve MERNİS tarama gibi işlevler. Kullanıcının sisteme giriş yaptıktan sonra karşılaştığı merkezi çalışma alanıdır.

Hasta yönetimi: Aktif, pasif, bekleyen, vefat etmiş ve silinmiş hasta listeleri; yeni kayıt, ilk kayıt, düzenleme, detaylı görüntüleme; birleşik hasta görünümü ve PDF dışa aktarma.

Ev ziyareti (izlem): Tamamlanan saha izlemleri, EK-3 konsültasyon formu, kaçırılan ziyaretler, izlem geçmişi.

Planlanan ziyaret: Gelecek ziyaret planları, çakışma kontrolü, pasif hasta plan temizliği.

Pansuman planı: Haftalık pansuman günleri planlama ve listeleme.

Günlük planlama: İlçe/bölge bazlı plan tablosu, yoğunluk kontrolü, skor kaydı.

İstatistikler: 80'den fazla rapor ekranı; KPI kartları, çapraz tablolar, veri sağlığı denetimi, e-Rapor uyum metrikleri, klinik karar desteği raporları.

Kullanıcı ve profil yönetimi: Personel CRUD, profil fotoğrafı, kurum ataması, kişisel istatistikler.
        """,
    )

    doc.add_heading("6.2. Operasyonel Modüller", level=2)
    add_paragraphs(
        doc,
        """
Branş randevu takvimi: Konsültasyon randevuları ve branş kotası yönetimi.

UHDS: Uzaktan sağlık hizmetleri randevu takvimi ve video görüşme (hasta ve personel tarafı).

Ekip planlama: Günlük ekip ve vardiya düzenlemesi.

Nöbet planı: Nöbet takvimi, izin/istek/tatil yönetimi, otomatik dağıtım.

Stok takibi: Malzeme tanımları, giriş-çıkış-iade, kritik stok uyarıları, hasta bazlı tüketim, sayım ve sipariş önerisi.

SMS bildirimleri: Toplu SMS, şablon yönetimi, günlük plan bildirimi, gönderim geçmişi.

Mesajlaşma: Personel arası DM, hasta konuşmaları, sistem duyuruları (broadcast).

e-Rapor: Elektronik rapor listesi, oluşturma ve hasta havuzu işlemleri.

İlaç ve tanı raporu: Hasta kartından ilaç ve tanı raporu kaydı.

İlaç rehberi: Etken madde ve ilaç arama (TİTCK hasta ilaçlarından bağımsız snapshot).

Hasta haritası: Aktif hastaların coğrafi dağılımı.

Adres koordinat bulma: Eksik koordinatlı adreslerin toplu eşleştirilmesi.
        """,
    )

    doc.add_heading("6.3. Yönetim Modülleri", level=2)
    add_paragraphs(
        doc,
        """
Kurum yönetimi, adres tanımları, branş/işlem/hastalık/güvence katalogları, araç listesi, personel izin/istek tanımları, hasta dosya sistemi (arşiv), tema yönetimi, uygulama ayarları, veritabanı bakımı, işlem günlüğü (denetim), ESYS/USBS uyum ve köprü modülleri, federasyon, REST API v1, KPS TC sorgusu ve CDN sürüm denetimi yönetim kapsamında sunulur.

Modül kayıt defteri (config/app-modules.registry.php) tüm modüllerin tek doğruluk kaynağıdır. Toggleable modüller ayar panelinden açılıp kapatılabilir; kurum ihtiyacına göre sistem sadeleştirilebilir veya genişletilebilir.
        """,
    )

    # --- 7. İSTATİSTİK ---
    doc.add_heading("7. İstatistik, Raporlama ve Veri Sağlığı", level=1)
    add_paragraphs(
        doc,
        """
SONEV'in istatistik modülü, yönetim kademesine operasyonel ve klinik karar desteği sağlayan kapsamlı bir raporlama hub'ıdır. Genel özet ve operasyon nabzı (operations pulse) ekranları birim performansının anlık görünümünü sunar. Demografik dağılımlar, güvence türü analizleri, hastalık profilleri ve yaş-cinsiyet bantları nüfus yapısını ortaya koyar.

İzlem sıklığı, personel iş yükü, bölgesel performans, aylık havuz hareketleri ve kayıt kohort analizleri operasyonel verimliliği ölçer. Barthel, Braden, MNA, Harizmi ve ITAKI dağılım raporları klinik risk profilini yansıtır. Veri sağlığı denetimi (data health) eksik TC, adres, doğum tarihi ve izlenmemiş aktif hasta gibi veri kalitesi sorunlarını tespit eder.

40'tan fazla çapraz tablo (xTab) raporu; yaş, ilçe, ziyaret durumu, personel, ay, bağımlılık düzeyi ve cihaz sayısı gibi boyutları çaprazlayarak derinlemesine analiz imkânı verir. e-Rapor hasta uyum metrikleri elektronik rapor süreçlerinin takibini kolaylaştırır. Klinik karar desteği raporu yüksek riskli ve izlem gecikmesi olan hastaları listeler.

Bu raporlama derinliği, ESYS'in ulusal kayıt ve bildirim odaklı yapısının ötesinde birim yöneticilerine yerel ve anlık karar desteği sağlar.
        """,
    )

    # --- 8. GÜVENLİK ---
    doc.add_heading("8. Güvenlik, KVKK ve Denetim Altyapısı", level=1)
    add_paragraphs(
        doc,
        """
Sağlık verisi hassas kişisel veri kapsamındadır. SONEV çok katmanlı güvenlik mimarisi uygular.

Kimlik doğrulama kullanıcı adı ve parola ile yapılır; isteğe bağlı e-imza (akıllı kart) ile giriş desteklenir. Oturum yönetimi güvenli oturum kimliği yenileme ve CSRF koruması içerir. Tüm POST isteklerinde CSRF token doğrulaması zorunludur; istemci tarafında formlara ve AJAX isteklerine otomatik token enjeksiyonu yapılır.

Yetkilendirme rol tabanlıdır; modül ve hasta düzeyinde erişim kontrolleri uygulanır (IDOR koruması). Kurum izolasyonu bir kurumun personelinin başka kurumun verisine erişmesini engeller.

HTTP güvenlik başlıkları (CSP, X-Frame-Options, Referrer-Policy, HSTS) yapılandırılmıştır. Giriş, kamu TC sorgusu ve geocode gibi hassas uç noktalarda hız sınırlama uygulanır. Yüklenen dosyalar doğrudan URL ile erişilemez; oturumlu proxy üzerinden yetki kontrolüyle sunulur.

İşlem günlüğü (AuditLog) hasta erişimi, izlem değişiklikleri, dışa aktarma, ESYS köprü işlemleri ve oturum açma/kapama olaylarını kayıt altına alır. Hasta portalı ve kamu sorgu ekranlarında KVKK bilgilendirme metinleri yer alır.
        """,
    )

    # --- 9. TEKNİK ---
    doc.add_heading("9. Teknik Mimari ve Altyapı", level=1)
    add_paragraphs(
        doc,
        """
SONEV, PHP 8 üzerinde çalışan özel bir MVC mimarisi kullanır. Giriş noktası public/index.php dosyasıdır; spl_autoload_register ile App\\ namespace altındaki sınıflar yüklenir. URL yönlendirmesi SEF yollar veya controller/action sorgu parametreleriyle çalışır.

Veritabanı erişimi PDO ile yapılır. Tablo öneki esh_ olarak kullanılır. Çoklu veritabanı desteği (MySQL, MSSQL, PostgreSQL, Oracle, SQLite) kurumsal IT altyapısı esnekliği sağlar.

Kullanıcı arayüzü Bootstrap 5 tabanlıdır. FormHelper standart CRUD ve filtre formlarında tutarlılık sağlar. Tema sistemi (default, meridian, evcare, aurora, winui vb.) kurumsal görünüm özelleştirmesine olanak tanır. İsteğe bağlı Vue 3 modern frontend pilotu dashboard özetini JSON API ile sunar.

Kalite güvencesi: PHP sözdizimi denetimi (php -l), PHPUnit smoke testleri, veritabanı migrasyon scriptleri ve tenant kapsam doğrulama aracı (verify_tenant_scope.php) mevcuttur.
        """,
    )

    # --- 10. ESYS KARŞILAŞTIRMA ---
    doc.add_heading("10. ESYS ile İlişki ve Karşılaştırmalı Değerlendirme", level=1)

    doc.add_heading("10.1. ESYS Nedir?", level=2)
    add_paragraphs(
        doc,
        """
Evde Sağlık Yönetim Sistemi (ESYS), Sağlık Bakanlığı tarafından yürütülen evde sağlık hizmetlerine ilişkin başvuru, değerlendirme, izlem, sevk, nakil ve taburculuk verilerinin kayıt altına alındığı ulusal dijital sistemdir. ESYS; AHBS, HBYS, UHDS ve e-Nabız gibi sistemlerle entegre çalışacak şekilde tasarlanmıştır. Başvurular ESKOM'a ve eş zamanlı olarak aile hekimliğine iletilir. ESYS, yönetmeliklerle tanımlanan resmî kayıt ve bildirim kanalıdır; birimlerin Bakanlık standartlarına uygun veri üretmesi ESYS üzerinden beklenir.

SONEV, ESYS'in yerine geçmek üzere değil; birim düzeyinde operasyonel mükemmelliği sağlarken ESYS ile uyumlu kalacak şekilde tasarlanmıştır. config/esys-field-mapping.json dosyasında hasta, başvuru, izlem, konsültasyon ve planlı izlem alanlarının ESYS karşılıkları tanımlıdır. esys_hasta_ref, esys_basvuru_ref, esys_izlem_ref gibi referans sütunları çift kayıt riskini azaltır. ESYS köprüsü JSON paketleri ile dışa/içe aktarma ve referans eşleme sağlar.
        """,
    )

    doc.add_heading("10.2. SONEV'in ESYS'e Kıyasla Gelişmiş ve Tamamlayıcı Yönleri", level=2)
    add_paragraphs(
        doc,
        """
Aşağıdaki tablo, ESYS'in ulusal koordinasyon rolü korunurken SONEV'in birim operasyonlarında sunduğu ek değeri özetler. ESYS zayıf veya kısıtlı olduğu alanlarda SONEV güçlü; ESYS'in zorunlu olduğu alanlarda SONEV köprü ve eşleme ile uyum sağlar.
        """,
    )

    add_table(
        doc,
        ["Alan", "ESYS (Ulusal sistem)", "SONEV (Birim operasyonel platform)"],
        [
            ["Temel rol", "Ulusal kayıt, bildirim, ESKOM/AHBS entegrasyonu", "Birim günlük operasyon, planlama, analiz"],
            ["Günlük rota planlama", "Sınırlı / standart plan kaydı", "TomTom rota optimizasyonu, mesafe matrisi, harita görselleştirme"],
            ["GPS saha doğrulama", "İzlem kaydında konum alanları (entegrasyonla)", "Check-in paneli, yapıldı izlemlerde GPS zorunluluğu seçeneği"],
            ["Klinik ölçekler", "Temel izlem verileri", "Barthel, Braden, Harizmi, ITAKI, MNA tam entegrasyonu ve trend"],
            ["Klinik bayraklar", "Standart hasta alanları", "20+ cihaz/bakım bayrağı, liste rozetleri, hızlı güncelleme"],
            ["Klinik karar desteği", "Yok", "Yüksek riskli + izlem gecikmesi otomatik tespit"],
            ["İstatistik derinliği", "Ulusal raporlama (Bakanlık düzeyi)", "80+ yerel rapor, 40+ çapraz tablo, veri sağlığı denetimi"],
            ["Stok yönetimi", "Yok", "Malzeme, giriş/çıkış, kritik stok, hasta tüketimi"],
            ["Ekip ve nöbet", "Yok", "Günlük ekip planı, nöbet otomatik dağıtım, izin/istek"],
            ["SMS ve mesajlaşma", "Ulusal kanallar (e-Nabız, ESHİM)", "Birim içi SMS şablonları, toplu bildirim, personel DM"],
            ["Hasta portalı", "e-Nabız üzerinden vatandaş erişimi", "TC+telefon ile plan/ziyaret özeti, SMS onayı"],
            ["Çoklu kurum", "İl/birim bazlı ESYS erişimi", "Tek kurulumda çok kurum, tenant izolasyonu, federasyon"],
            ["Özelleştirme", "Bakanlık standartları (sabit)", "Modüler aç/kapa, tema, kurumsal ayarlar JSON"],
            ["Veri egemenliği", "Merkezî bulut (Bakanlık)", "On-premise, kurum sunucusunda tam veri kontrolü"],
            ["Offline / yerel ağ", "İnternet bağımlı", "Yerel sunucuda çalışır; internet kesintisinde operasyon devam eder"],
            ["REST API", "Resmî API (kısıtlı/merkezî)", "Bearer token ile patients/visits/plans JSON (isteğe bağlı)"],
            ["Denetim günlüğü", "Ulusal sistem logları", "Kurum içi KVKK denetim günlüğü, dışa aktarma izi"],
            ["e-imza girişi", "Kamu PKI entegrasyonu", "Yerel akıllı kart köprüsü, sertifika pinning"],
        ],
    )

    doc.add_heading("10.3. ESYS Köprüsü ve Uyum Stratejisi", level=2)
    add_paragraphs(
        doc,
        """
SONEV'in ESYS stratejisi «önce operasyonel mükemmellik, paralelde ulusal uyum» ilkesine dayanır:

1. Alan eşlemesi: config/esys-field-mapping.json ile ESH alanları ESYS karşılıklarına haritalanır (hasta kimlik, adres, tanı, izlem, konsültasyon, planlı izlem).

2. Referans numaraları: esys_hasta_ref, esys_basvuru_ref, esys_izlem_ref, esys_konsultasyon_ref, esys_plan_ref sütunları çift kayıt yönetimini kolaylaştırır.

3. Dosya köprüsü: JSON paket dışa/içe aktarma; eksik referanslı hastaları filtreleme; izlem gün aralığı seçimi; senkronizasyon günlüğü (esh_esys_sync_log).

4. API hazırlığı: Gelecekteki HSYS/ESYS HTTP köprüsü için api_base_url ve api_enabled ayarları mevcuttur.

5. USBS/e-Nabız köprüsü: İzlem bildirimi için ayrı USBS uyum modülü ve köprüsü.

Bu yapı sayesinde birim, günlük işini SONEV'de hızlı ve verimli yürütür; Bakanlık bildirim gereksinimleri için ESYS ile senkronize veri paketleri üretebilir. Resmî entegrasyon API'leri yayınlandığında SONEV köprüsü API moduna geçirilebilir.
        """,
    )

    doc.add_heading("10.4. Birlikte Kullanım Senaryosu", level=2)
    add_paragraphs(
        doc,
        """
Önerilen hibrit model: SONEV günlük operasyonel omurga olarak kullanılır; ESYS ulusal kayıt ve bildirim kanalı olarak korunur. Hasta ilk kaydı SONEV'de yapılır, klinik profil ve planlar detaylı işlenir. Periyodik olarak ESYS köprüsü ile hasta ve izlem verileri JSON paket olarak dışa aktarılır veya ESYS referans numaraları manuel/otomatik eşlenir. Bakanlık raporları ESYS üzerinden verilir; birim içi performans ve kalite raporları SONEV istatistik modülünden alınır.

Bu model, birçok ilde ESYS'e ek olarak yerel Excel, kağıt veya eski yazılımlarla yürütülen operasyonları tek platformda toplarken ulusal uyumu korur.
        """,
    )

    # --- 11. ULUSLARARASI KARŞILAŞTIRMA ---
    doc.add_heading("11. Dünyadaki Benzer Sistemlerle Karşılaştırma", level=1)

    doc.add_heading("11.1. Küresel Pazar Bağlamı", level=2)
    add_paragraphs(
        doc,
        """
Evde sağlık ve evde bakım yazılımları dünya genelinde hızla büyüyen bir pazar segmentidir. ABD'de WellSky (eski adıyla Kinnser), Axxess ve AlayaCare; Kanada'da AlayaCare ve CellTrak; Avustralya'da Clever Care ve AlayaCare; İngiltere'de EMIS Community ve SystmOne Community; Almanya'da Medifox ve Snap Core gibi çözümler yaygındır. Büyük hastane bilgi sistemleri (Epic, Cerner/Oracle Health) de home health modülleri sunar.

Bu sistemlerin ortak özellikleri: hasta planlama, ziyaret belgeleme (visit documentation), personel zamanlama, faturalandırma (billing) ve uyum raporlamasıdır. Ancak ülke sağlık sistemine özgü entegrasyonlar, fiyatlandırma modeli ve veri egemenliği açısından önemli farklar bulunur.
        """,
    )

    doc.add_heading("11.2. SONEV'in Uluslararası Benzerlere Göre Artıları", level=2)

    add_table(
        doc,
        ["Özellik", "Tipik global çözüm", "SONEV avantajı"],
        [
            ["Fiyatlandırma", "SaaS abonelik, kullanıcı/visit başına ücret", "On-premise, lisans/abonelik bağımlılığı yok; kurum sunucusunda"],
            ["Türkiye uyumu", "Genel ICD-10; yerel entegrasyon ek maliyet", "ESYS/AHBS/USBS alan eşlemesi yerleşik; EK-3, e-Rapor, TC kimlik"],
            ["Çoklu kurum", "Çoğunlukla tek organizasyon veya ayrı tenant başına ücret", "Tek kurulumda çok kurum; federasyon modülü"],
            ["Rota optimizasyonu", "Premium eklenti veya üçüncü parti", "TomTom entegrasyonu yerleşik; rota cache"],
            ["Klinik ölçekler", "Sınırlı veya ayrı modül", "5 ölçek + klinik karar desteği dahili"],
            ["Stok", "Ayrı ERP entegrasyonu gerekir", "Dahili stok modülü, hasta tüketim bağlantısı"],
            ["Nöbet / ekip", "Workforce modülü ek ücret", "Nöbet otomatik dağıtım, ekip planı dahili"],
            ["İstatistik", "Standart dashboard", "80+ rapor, 40+ çapraz tablo, veri sağlığı"],
            ["Mesajlaşma / SMS", "Üçüncü parti", "Dahili mesajlaşma ve SMS şablonları"],
            ["Hasta portalı", "Ayrı uygulama", "Dahili portal (TC+telefon)"],
            ["e-imza", "Ülkeye özgü değil", "Türkiye akıllı kart e-imza girişi"],
            ["Teknik bağımlılık", "Cloud vendor lock-in", "Düz PHP 8, PDO; Composer/npm opsiyonel"],
            ["Veritabanı", "Genellikle vendor cloud DB", "MySQL, MSSQL, PostgreSQL, Oracle, SQLite seçeneği"],
            ["Özelleştirme", "Konfigürasyon sınırlı", "Modül registry, tema, kurumsal JSON ayarlar"],
            ["Faturalandırma", "ABD odaklı (Medicare/Medicaid)", "Türkiye güvence modeli; faturalandırma yerine klinik/operasyon odaklı"],
        ],
    )

    doc.add_heading("11.3. WellSky / AlayaCare Tipi Platformlarla Detaylı Kıyas", level=2)
    add_paragraphs(
        doc,
        """
WellSky (Kinnser) ABD evde sağlık pazarının liderlerinden biridir. OASIS değerlendirmeleri, Medicare uyumu, elektronik imza ve mobil saha uygulaması güçlüdür. Ancak Türkiye sağlık finansmanı, ESYS bildirimi, EK-3 konsültasyon formu, e-Rapor ve TC kimlik doğrulama gibi yerel gereksinimlere doğrudan uyumlu değildir. SONEV bu yerel gereksinimleri yerleşik karşılar.

AlayaCare kanada kökenli, bulut tabanlı evde bakım platformudur. Zamanlama, rota, mobil uygulama ve aile portalı güçlüdür. Ancak SaaS modeli veri egemenliği endişeleri yaratabilir; Türkiye'deki kamu hastaneleri ve belediye birimleri için on-premise SONEV daha uygun bir profil sunar. AlayaCare'in faturalandırma ve CRM odaklı yapısı, Türkiye evde sağlık birimlerinin klinik ve operasyonel önceliklerinden farklılaşır.

Epic ve Cerner home health modülleri büyük sağlık sistemlerine entegre çalışır; ancak lisans maliyeti, implementasyon süresi ve Türkiye yerelleştirmesi açısından evde sağlık birimi ölçeğinde ağır kalır. SONEV hafif, hızlı kurulabilir ve Türkiye bağlamına özgüdür.
        """,
    )

    doc.add_heading("11.4. SONEV'in Ayırt Edici Değer Önerisi", level=2)
    add_paragraphs(
        doc,
        """
SONEV'i küresel rakiplerinden ayıran temel değer önerileri:

1. Türkiye sağlık ekosistemi yerleşimi: ESYS, AHBS, USBS, e-Rapor, EK-3, KPS, TC kimlik ve güvence modeli.

2. Operasyonel derinlik: Rota, GPS check-in, stok, nöbet, ekip, SMS, mesajlaşma — tek platformda.

3. Klinik zenginlik: Ölçekler, bayraklar, karar desteği, pansuman, yara fotoğrafı.

4. Analitik güç: Kurum içi 80+ rapor; Bakanlık raporlarından bağımsız yerel KPI.

5. Veri egemenliği: On-premise; internet kesintisinde yerel operasyon.

6. Maliyet etkinliği: SaaS abonelik yok; mevcut XAMPP/Apache + MariaDB yeterli.

7. Modüler esneklik: Küçük birim sade mod; büyük kurum tam mod.

8. Çok kurum: İl genelinde tek platform, kurum bazlı izolasyon.

9. Açık köprü mimarisi: ESYS/USBS JSON köprüleri; REST API; federasyon.

10. Düşük teknik borç riski: Düz PHP, anlaşılır MVC; vendor lock-in yok.
        """,
    )

    # --- 12. ÇOKLU KURUM ---
    doc.add_heading("12. Çoklu Kurum Mimarisi", level=1)
    add_paragraphs(
        doc,
        """
SONEV 4.0 ile çoklu kurum (multi-tenant) desteği kazanmıştır. Tek veritabanı ve tek uygulama örneği üzerinde birden fazla evde sağlık birimi bağımsız çalışabilir. Her hasta kaydı bir kuruma bağlıdır; (kurum_id, tckimlik) birlikte benzersizdir.

TenantContext ve TenantSqlHelper oturum bazlı kurum filtresi uygular. Kurumsal ayarlar esh_kurumlar.ayarlar_json alanında saklanır. Platform kataloğu (branş, işlem, istek) kurum_id=0 ile süper yönetici tarafından yönetilir; kurumlar junction tabloları (esh_kurum_brans, esh_kurum_istek, esh_kurum_islem) üzerinden seçim yapar.

Federasyon modülü bölgeler arası veri köprüsü ve çok bölgeli SaaS senaryoları için export/import bundle desteği sunar. Bu mimari, il sağlık müdürlüğü düzeyinde merkezi platform + ilçe birimleri modeline uygundur.
        """,
    )

    # --- 13. KURULUM ---
    doc.add_heading("13. Kurulum, İşletim ve Sürdürülebilirlik", level=1)
    add_paragraphs(
        doc,
        """
Minimum gereksinimler: PHP 8.x, MariaDB/MySQL, web sunucusu (Apache önerilir). Uygulama public/ dizini altından servis edilir. İlk kurulum web tabanlı kurulum sihirbazı ile yapılabilir.

Üretim ortamında düzenli veritabanı yedeklemesi önerilir. Veritabanı bakımı modülü tam yedek, tablo yedek, geri yükleme ve optimizasyon sunar. Güncellemeler migrasyon scriptleri ve changelog ile takip edilir. Smoke testler (php tools/run_smoke_tests.php) temel regresyon kontrolü sağlar.

İsteğe bağlı entegrasyonlar: TomTom harita anahtarı, SMS gateway, e-imza akıllı kart köprüsü, KPS bağlantısı. Bunlar çekirdek iş akışları için zorunlu değildir.
        """,
    )

    # --- 14. SONUÇ ---
    doc.add_heading("14. Sonuç ve Stratejik Değerlendirme", level=1)
    add_paragraphs(
        doc,
        """
SONEV, Türkiye'deki evde sağlık birimleri için tasarlanmış, operasyonel derinliği yüksek, modüler ve çok kurumlu bir bilgi yönetim sistemidir. ESYS ulusal koordinasyon ve bildirim kanalı olarak kalırken, SONEV birim düzeyinde günlük operasyonların dijital omurgasını oluşturur.

ESYS'e kıyasla SONEV; rota optimizasyonu, GPS saha doğrulama, klinik ölçekler, klinik karar desteği, kapsamlı istatistik, stok, nöbet, ekip planı, SMS, mesajlaşma, hasta portalı, çoklu kurum ve on-premise veri egemenliği gibi alanlarda belirgin üstünlükler sunar. ESYS köprüsü ve alan eşlemesi ile ulusal uyum korunur.

Dünya genelindeki WellSky, AlayaCare, Epic Home Health gibi platformlara kıyasla SONEV; Türkiye sağlık mevzuatına yerleşik uyum, maliyet etkinliği, veri egemenliği, modüler esneklik ve operasyonel-klinik derinlik kombinasyonu ile ayırt edilir.

SONEV ile evde sağlık birimleri; hasta dosyasını merkezileştirir, saha operasyonlarını dijitalleştirir, veriye dayalı yönetim kararları alır, KVKK ve denetim gereksinimlerini karşılar ve ulusal sistemlerle köprülü çalışarak hem yerel verimlilik hem ulusal uyum hedeflerine ulaşır.
        """,
    )

    # --- EK ---
    doc.add_page_break()
    doc.add_heading("Ek A: Modül Listesi (Özet)", level=1)

    modules = [
        ("dashboard", "Ana panel", "Çekirdek"),
        ("patient", "Hasta yönetimi", "Çekirdek"),
        ("visit", "Ev ziyareti (izlem)", "Çekirdek"),
        ("planned_visit", "Planlanan ziyaret", "Çekirdek"),
        ("pansuman", "Pansuman planı", "Çekirdek"),
        ("planning", "Günlük planlama", "Çekirdek"),
        ("stats", "İstatistikler", "Çekirdek"),
        ("erapor", "e-Rapor", "Site"),
        ("randevu", "Branş randevu takvimi", "Site"),
        ("uhds", "UHDS", "Site"),
        ("stok", "Stok takibi", "Site"),
        ("sms_bildirim", "SMS bildirimleri", "Site"),
        ("mesajlasma", "Mesajlaşma", "Site"),
        ("harita", "Hasta haritası", "Yönetim"),
        ("nobet", "Nöbet planı", "Yönetim"),
        ("ekip", "Ekip planlama", "Yönetim"),
        ("esys_compliance", "ESYS uyum", "Yönetim"),
        ("usbs_compliance", "USBS uyum", "Yönetim"),
        ("federation", "Federasyon", "Yönetim"),
        ("rest_api", "REST API v1", "Yönetim"),
        ("patient_portal", "Hasta portalı", "Kamu"),
    ]
    add_table(doc, ["Anahtar", "Modül", "Grup"], modules)

    doc.add_heading("Ek B: Sürüm ve Teknik Bilgiler", level=1)
    add_paragraphs(
        doc,
        """
Uygulama adı: SONEV (ESH)
Sürüm: 4.0.0
Kod adı: Burutay
Yayın tarihi: 2026-06-17
Backend: PHP 8, özel MVC
Veritabanı: MariaDB/MySQL (varsayılan)
Controller sayısı: 55+
Modül sayısı: 40+
İstatistik ekranı: 80+
        """,
    )

    doc.add_heading("Ek C: Yasal Uyarı", level=1)
    add_paragraphs(
        doc,
        """
Bu belgedeki ESYS ve uluslararası yazılım karşılaştırmaları genel bilgilendirme amaçlıdır. ESYS, Sağlık Bakanlığı'nın resmî sistemidir; bildirim ve kayıt yükümlülükleri ilgili mevzuata tabidir. SONEV'in ESYS ile kullanımı kurumun yasal süreçlerine uygun şekilde planlanmalıdır. Uluslararası ürün karşılaştırmaları kamuya açık ürün bilgilerine dayanır; ticari ürünlerin özellikleri zaman içinde değişebilir.
        """,
    )

    return doc


def main() -> int:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"Olusturuldu: {OUTPUT}")
    print(f"Boyut: {OUTPUT.stat().st_size:,} byte")
    return 0


if __name__ == "__main__":
    sys.exit(main())
